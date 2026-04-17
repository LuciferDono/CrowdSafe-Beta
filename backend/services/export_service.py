"""
Export service for analytics data.
Generates CSV, DOCX, PDF, and Markdown reports from metrics data.
"""

import csv
import io
from datetime import datetime, timezone, timedelta

IST = timezone(timedelta(hours=5, minutes=30))


def _to_ist(ts_str):
    """Convert ISO timestamp string to IST formatted string."""
    if not ts_str:
        return ''
    try:
        dt = datetime.fromisoformat(str(ts_str).replace('Z', '+00:00'))
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(IST).strftime('%d %b %Y, %I:%M:%S %p IST')
    except (ValueError, TypeError):
        return str(ts_str)


def _now_ist():
    return datetime.now(IST).strftime('%d %b %Y, %I:%M:%S %p IST')


def export_csv(metrics, summary):
    """Generate CSV bytes from metrics list."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        'Timestamp', 'Count', 'Density (p/m²)', 'Avg Velocity (m/s)',
        'Max Velocity (m/s)', 'Surge Rate', 'Risk Score', 'Risk Level',
        'Capacity Util (%)', 'Flow In', 'Flow Out'
    ])
    for m in metrics:
        writer.writerow([
            _to_ist(m.get('timestamp', '')),
            m.get('count', 0),
            m.get('density', 0),
            m.get('avg_velocity', 0),
            m.get('max_velocity', 0),
            m.get('surge_rate', 0),
            m.get('risk_score', 0),
            m.get('risk_level', 'SAFE'),
            m.get('capacity_utilization', 0),
            m.get('flow_in', 0),
            m.get('flow_out', 0),
        ])
    return output.getvalue().encode('utf-8')


def export_docx(metrics, summary, camera_name='Camera'):
    """Generate DOCX bytes from metrics data."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading('CrowdSafe Analytics Report', level=0)
    doc.add_paragraph(f'Camera: {camera_name}')
    doc.add_paragraph(f'Generated: {_now_ist()}')
    doc.add_paragraph(f'Records: {len(metrics)}')

    # Summary section
    doc.add_heading('Summary', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = summary_table.rows
    rows[0].cells[0].text = 'Average People Count'
    rows[0].cells[1].text = str(round(summary.get('avg_count', 0), 1))
    rows[1].cells[0].text = 'Peak Count'
    rows[1].cells[1].text = str(summary.get('peak_count', 0))
    rows[2].cells[0].text = 'Average Density (p/m²)'
    rows[2].cells[1].text = str(round(summary.get('avg_density', 0), 3))
    rows[3].cells[0].text = 'Max Risk Score'
    rows[3].cells[1].text = f"{round((summary.get('max_risk_score', 0)) * 100, 1)}%"

    # Data table
    doc.add_heading('Metrics Data', level=1)
    headers = ['Timestamp', 'Count', 'Density', 'Velocity', 'Risk Score', 'Risk Level']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for m in metrics:
        row = table.add_row()
        row.cells[0].text = _to_ist(m.get('timestamp', ''))
        row.cells[1].text = str(m.get('count', 0))
        row.cells[2].text = str(m.get('density', 0))
        row.cells[3].text = str(m.get('avg_velocity', 0))
        row.cells[4].text = f"{round(m.get('risk_score', 0) * 100, 1)}%"
        row.cells[5].text = m.get('risk_level', 'SAFE')

        # Font size for data
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(metrics, summary, camera_name='Camera'):
    """Generate PDF bytes from metrics data."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font('Helvetica', 'B', 18)
    pdf.cell(0, 12, 'CrowdSafe Analytics Report', new_x='LMARGIN', new_y='NEXT', align='C')
    pdf.ln(4)

    # Metadata
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 6, f'Camera: {camera_name}', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(0, 6, f'Generated: {_now_ist()}', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(0, 6, f'Total Records: {len(metrics)}', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(6)

    # Summary box
    pdf.set_font('Helvetica', 'B', 13)
    pdf.cell(0, 8, 'Summary', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(2)

    pdf.set_font('Helvetica', '', 10)
    summary_data = [
        ('Average People Count', str(round(summary.get('avg_count', 0), 1))),
        ('Peak Count', str(summary.get('peak_count', 0))),
        ('Average Density (p/m²)', str(round(summary.get('avg_density', 0), 3))),
        ('Max Risk Score', f"{round(summary.get('max_risk_score', 0) * 100, 1)}%"),
    ]
    for label, val in summary_data:
        pdf.set_font('Helvetica', 'B', 10)
        pdf.cell(80, 7, label, border=1)
        pdf.set_font('Helvetica', '', 10)
        pdf.cell(0, 7, val, border=1, new_x='LMARGIN', new_y='NEXT')
    pdf.ln(6)

    # Data table
    pdf.set_font('Helvetica', 'B', 13)
    pdf.cell(0, 8, 'Metrics Data', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(2)

    headers = ['Timestamp (IST)', 'Count', 'Density', 'Velocity', 'Risk %', 'Level']
    col_widths = [58, 16, 20, 20, 18, 20]

    # Header
    pdf.set_font('Helvetica', 'B', 8)
    pdf.set_fill_color(60, 60, 80)
    pdf.set_text_color(255, 255, 255)
    for i, h in enumerate(headers):
        pdf.cell(col_widths[i], 7, h, border=1, fill=True, align='C')
    pdf.ln()

    # Data
    pdf.set_font('Helvetica', '', 7)
    pdf.set_text_color(0, 0, 0)
    for m in metrics:
        ts = _to_ist(m.get('timestamp', ''))

        risk_pct = f"{round(m.get('risk_score', 0) * 100, 1)}"
        level = m.get('risk_level', 'SAFE')

        row_data = [
            str(ts),
            str(m.get('count', 0)),
            str(m.get('density', 0)),
            str(m.get('avg_velocity', 0)),
            risk_pct,
            level,
        ]

        # Color-code risk level
        if level == 'CRITICAL':
            pdf.set_fill_color(255, 200, 200)
        elif level == 'WARNING':
            pdf.set_fill_color(255, 235, 200)
        elif level == 'CAUTION':
            pdf.set_fill_color(255, 255, 200)
        else:
            pdf.set_fill_color(255, 255, 255)

        for i, val in enumerate(row_data):
            pdf.cell(col_widths[i], 6, val, border=1, fill=True, align='C')
        pdf.ln()

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def export_markdown(metrics, summary, camera_name='Camera'):
    """Generate Markdown string from metrics data."""
    lines = []
    lines.append('# CrowdSafe Analytics Report')
    lines.append('')
    lines.append(f'**Camera:** {camera_name}')
    lines.append(f'**Generated:** {_now_ist()}')
    lines.append(f'**Records:** {len(metrics)}')
    lines.append('')
    lines.append('## Summary')
    lines.append('')
    lines.append('| Metric | Value |')
    lines.append('|---|---|')
    lines.append(f'| Average People Count | {round(summary.get("avg_count", 0), 1)} |')
    lines.append(f'| Peak Count | {summary.get("peak_count", 0)} |')
    lines.append(f'| Average Density (p/m²) | {round(summary.get("avg_density", 0), 3)} |')
    lines.append(f'| Max Risk Score | {round(summary.get("max_risk_score", 0) * 100, 1)}% |')
    lines.append('')
    lines.append('## Metrics Data')
    lines.append('')
    lines.append('| Timestamp | Count | Density | Velocity | Risk Score | Risk Level |')
    lines.append('|---|---|---|---|---|---|')

    for m in metrics:
        ts = _to_ist(m.get('timestamp', ''))
        risk_pct = f"{round(m.get('risk_score', 0) * 100, 1)}%"
        lines.append(
            f'| {ts} | {m.get("count", 0)} | {m.get("density", 0)} '
            f'| {m.get("avg_velocity", 0)} | {risk_pct} | {m.get("risk_level", "SAFE")} |'
        )

    lines.append('')
    return '\n'.join(lines)
